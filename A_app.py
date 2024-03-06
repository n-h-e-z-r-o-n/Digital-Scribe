

from tkinter import filedialog

"""
import docx

#file_path = filedialog.askopenfilename()
#print(file_path)

document = docx.Document('C:\HEZRON WEKESA\Downloads\CSC 451 - Distributed Database by Dr. Kahonge.docx')
# Do something with the document, such as printing its content
for paragraph in document.paragraphs:
    print(paragraph.text)
"""

import tkinter as tk
from tkinter import filedialog


def open_file_dialog():
    root = tk.Tk()
    root.withdraw()  # Hide the main window

    document = docx.Document('C:\HEZRON WEKESA\Downloads\CSC 451 - Distributed Database by Dr. Kahonge.docx')

    file_path = filedialog.askopenfilename()

    if file_path:
        print("Selected file path:", file_path)
    else:
        print("No file selected")


# Call the function to open the file dialog
open_file_dialog()


